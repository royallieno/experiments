import os
import time
from datetime import datetime
import glob
import sys

def check_virtual_env():
    """Check if running in virtual environment."""
    if not hasattr(sys, 'real_prefix') and not (hasattr(sys, 'base_prefix') and sys.base_prefix != sys.prefix):
        print("Error: Script must be run in virtual environment!")
        print("\nTo set up and use virtual environment:")
        print("1. Run setup script first:")
        print("   bash setup.sh")
        print("\n2. Activate virtual environment:")
        print("   source venv/bin/activate")
        print("\n3. Then run this script again:")
        print("   python doc_rephraser.py")
        sys.exit(1)

def read_document(file_path):
    """Read content from document while preserving structure."""
    file_extension = os.path.splitext(file_path)[1].lower()
    
    if file_extension == '.docx':
        try:
            from docx import Document
            doc = Document(file_path)
            # Return document object and paragraphs
            return doc, [paragraph for paragraph in doc.paragraphs if paragraph.text.strip()]
        except Exception as e:
            print(f"Error reading .docx file: {e}")
            return None, None
    else:
        print(f"Unsupported file format: {file_extension}. Please use .docx files.")
        return None, None

def rephrase_text(text, rephraser):
    """Rephrases the given text while maintaining meaning but with different wording."""
    try:
        if not text.strip():
            return text
            
        # Use the rephraser with parameters optimized for paraphrasing
        rephrased = rephraser(
            text,
            max_length=int(len(text.split()) * 1.2),  # Allow slight expansion
            min_length=int(len(text.split()) * 0.8),  # Prevent excessive compression
            num_beams=5,  # Increase beam search for better alternatives
            do_sample=True,  # Enable sampling for varied phrasing
            temperature=0.6,  # Lower temperature for more focused rephrasing
            top_k=50,  # Limit vocabulary choices for more coherent output
            top_p=0.95,  # High nucleus sampling for quality
            repetition_penalty=2.5,  # Avoid repetitive phrases
            length_penalty=1.0,  # Balanced length control
            no_repeat_ngram_size=3  # Prevent repetition of phrases
        )[0]['summary_text']
        
        return rephrased.strip()
    except Exception as e:
        print(f"Error rephrasing text: {e}")
        return text

def process_file(input_file_path, rephraser):
    """Process a single file."""
    input_filename = os.path.basename(input_file_path)
    print(f"Processing {input_filename}...")
    
    try:
        # Read document content
        doc, paragraphs = read_document(input_file_path)
        if doc is None or paragraphs is None:
            return False
        
        # Create new document with same style
        from docx import Document
        new_doc = Document()
        
        # Copy document styles and settings
        for style in doc.styles:
            if style.name not in new_doc.styles:
                try:
                    new_doc.styles.add_style(style.name, style.type)
                except:
                    pass

        # Copy sections and their properties
        for section in doc.sections:
            new_section = new_doc.sections[-1]  # Get the last section
            new_section.page_height = section.page_height
            new_section.page_width = section.page_width
            new_section.left_margin = section.left_margin
            new_section.right_margin = section.right_margin
            new_section.top_margin = section.top_margin
            new_section.bottom_margin = section.bottom_margin
            new_section.header_distance = section.header_distance
            new_section.footer_distance = section.footer_distance

        # Process each paragraph
        total_paragraphs = len(paragraphs)
        for i, para in enumerate(paragraphs, 1):
            print(f"Processing paragraph {i}/{total_paragraphs}...")
            
            # Skip empty paragraphs
            if not para.text.strip():
                continue
                
            # Create new paragraph with same style and alignment
            new_para = new_doc.add_paragraph()
            new_para.style = para.style
            new_para.alignment = para.alignment
            
            # Rephrase the paragraph text
            rephrased_text = rephrase_text(para.text, rephraser)
            
            # Split rephrased text to match original runs
            words = rephrased_text.split()
            word_index = 0
            
            # Preserve formatting from original runs
            for run in para.runs:
                if not words[word_index:]:
                    break
                    
                # Calculate approximate number of words for this run
                run_word_count = len(run.text.split())
                run_words = ' '.join(words[word_index:word_index + run_word_count])
                word_index += run_word_count
                
                # Create new run with preserved formatting
                new_run = new_para.add_run(run_words + ' ')
                
                # Copy run formatting
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                new_run.font.name = run.font.name
                if run.font.size:
                    new_run.font.size = run.font.size
                if run.font.color.rgb:
                    new_run.font.color.rgb = run.font.color.rgb
                
            # Add a small delay to avoid rate limiting
            time.sleep(1)
        
        # Generate output filename with timestamp
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"rephrased_{input_filename.split('.')[0]}_{timestamp}.docx"
        output_path = os.path.join('output', output_filename)
        
        # Save the rephrased document
        new_doc.save(output_path)
            
        print(f"Document {input_filename} rephrased successfully. Output saved to: {output_path}")
        return True
        
    except Exception as e:
        print(f"Error processing {input_filename}: {e}")
        return False

def process_documents():
    """Process all documents in the input directory."""
    try:
        # Check if running in virtual environment
        check_virtual_env()

        # Import after virtual environment check
        try:
            from transformers import pipeline
        except ImportError:
            print("Error: Required packages not installed!")
            print("Please run 'bash setup.sh' first")
            return
        
        # Initialize the rephraser with a model better suited for paraphrasing
        print("Loading language model... (this might take a few moments)")
        rephraser = pipeline(
            "summarization",
            model="facebook/bart-large-cnn",  # Changed to a model better at maintaining details
            device=-1,  # Force CPU usage
            framework="pt",
            max_length=1024,
            min_length=8
        )
        
        # Get all files from input directory
        input_dir = 'input'
        if not os.path.exists(input_dir):
            print(f"Input directory '{input_dir}' not found. Creating it...")
            os.makedirs(input_dir)
            print(f"Please place your .docx files in the '{input_dir}' directory and run the script again.")
            return
            
        input_files = glob.glob(os.path.join(input_dir, '*.docx'))  # Only process .docx files
        
        if not input_files:
            print(f"No .docx files found in '{input_dir}' directory")
            print(f"Please place your .docx files in the '{input_dir}' directory and run the script again.")
            return
        
        print(f"Found {len(input_files)} .docx files to process:")
        for f in input_files:
            print(f"- {os.path.basename(f)}")
        
        # Create output folder if it doesn't exist
        if not os.path.exists('output'):
            os.makedirs('output')
        
        # Process each file
        successful = 0
        failed = 0
        for input_file_path in input_files:
            if process_file(input_file_path, rephraser):
                successful += 1
            else:
                failed += 1
        
        # Print summary
        print(f"\nProcessing complete!")
        print(f"Successfully processed: {successful} files")
        if failed > 0:
            print(f"Failed to process: {failed} files")
                
    except Exception as e:
        print(f"An error occurred: {e}")

if __name__ == "__main__":
    process_documents() 